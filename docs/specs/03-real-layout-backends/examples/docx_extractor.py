"""ILLUSTRATIVE — proposed DocxExtractor for app/layout.py. NOT wired in.

Body via python-docx's public API; footnotes via the word/footnotes.xml zip part
(python-docx does NOT expose footnotes -- naive use silently drops them).

Two traps this handles (spec 03 §2 of design.md, both acceptance criteria):
  1. Word ALWAYS writes <w:footnote w:type="separator" w:id="-1"> and
     <w:footnote w:type="continuationSeparator" w:id="0">. Skip any w:footnote
     with a w:type OR id <= 0, else every document grows spurious footnote blocks.
  2. w:footnoteReference/@w:id resolves to w:footnote/@w:id (footnote-local ids,
     independent of endnotes.xml). Endnotes are out of scope for v1.

Real seam (app/layout.py): register("docx", DocxExtractor()); extract -> Document.
Canonical types from spec 02 (imported from app.layout per its seam).

Do NOT import this file. Port the pieces into app/layout.py.
"""
from __future__ import annotations

import io
import zipfile

from docx import Document as DocxDocument

from app.layout import SCHEMA_VERSION, Block, Document, Relation  # spec-02 types

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _heading_level(style_name: str) -> int:
    # "Heading 1".."Heading 9" -> 1..9; "Title" -> 1; else 0 (body paragraph).
    if style_name and style_name.startswith("Heading "):
        tail = style_name.removeprefix("Heading ").strip()
        return int(tail) if tail.isdigit() else 0
    return 1 if style_name == "Title" else 0


class DocxExtractor:
    def extract(self, raw: str | bytes, source: str) -> Document:
        if isinstance(raw, str):
            raise TypeError("DocxExtractor expects raw bytes, not str")
        doc = DocxDocument(io.BytesIO(raw))

        blocks: list[Block] = []
        relations: list[Relation] = []
        bid = 0

        # --- footnotes from the zip part (real footnotes only) ---
        fn_blocks: dict[str, Block] = {}  # w:id -> footnote Block
        for fid, text in self._read_footnotes(raw).items():
            bid += 1
            blk = Block(id=f"b{bid}", type="footnote", text=text, attrs={"w_id": fid})
            blocks.append(blk)
            fn_blocks[fid] = blk

        # --- body paragraphs + tables, in document order ---
        for para in doc.paragraphs:
            text = para.text.strip()
            cited = self._footnote_refs(para)  # w:ids cited in this paragraph
            if not text and not cited:
                continue
            bid += 1
            level = _heading_level(para.style.name if para.style else "")
            blk = Block(id=f"b{bid}", type="heading" if level else "paragraph",
                        text=text, level=level)
            blocks.append(blk)
            for fid in cited:
                if fid in fn_blocks:
                    relations.append(Relation("footnote_ref", blk.id, fn_blocks[fid].id))

        for table in doc.tables:
            rows = ["\t".join(c.text.strip() for c in r.cells) for r in table.rows]
            text = "\n".join(rows).strip()
            if not text:
                continue
            bid += 1
            blocks.append(Block(id=f"b{bid}", type="table", text=text))

        return Document(schema_version=SCHEMA_VERSION, source=source, format="docx",
                        meta={}, blocks=blocks, relations=relations)

    # --- footnote part access (not on the public API) ----------------------

    def _read_footnotes(self, raw: bytes) -> dict[str, str]:
        out: dict[str, str] = {}
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            if "word/footnotes.xml" not in zf.namelist():
                return out
            import xml.etree.ElementTree as ET
            root = ET.fromstring(zf.read("word/footnotes.xml"))
            for fn in root.findall(f"{_W}footnote"):
                ftype = fn.get(f"{_W}type")
                fid = fn.get(f"{_W}id")
                # skip separator / continuationSeparator pseudo-footnotes
                if ftype is not None or fid is None or int(fid) <= 0:
                    continue
                text = "".join(t.text or "" for t in fn.iter(f"{_W}t")).strip()
                if text:
                    out[fid] = text
        return out

    def _footnote_refs(self, para) -> list[str]:
        # python-docx exposes the lxml element as para._p; refs are not on the API.
        ns = {"w": _W.strip("{}")}
        return [r.get(f"{_W}id") for r in para._p.findall(".//w:footnoteReference", ns)]
